#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Конвертер Markdown в Word DOCX
Конвертирует TZ_NowPlay_FINAL.md в TZ_NowPlay_FINAL.docx
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def parse_markdown_to_docx(md_file, docx_file):
    """Парсит Markdown файл и создает DOCX документ"""
    
    # Создаем новый документ
    doc = Document()
    
    # Настраиваем стили
    setup_styles(doc)
    
    # Читаем Markdown файл
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Разбиваем на строки
    lines = content.split('\n')
    
    i = 0
    in_table = False
    table_data = []
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Пропускаем HTML комментарии и теги
        if line.strip().startswith('<') or line.strip().startswith('<!--'):
            i += 1
            continue
        
        # Обработка блоков кода
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # Конец блока кода
                in_code_block = False
                if code_lines:
                    p = doc.add_paragraph('\n'.join(code_lines))
                    p.style = 'Code'
                code_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Обработка таблиц
        if '|' in line and not line.strip().startswith('#'):
            if not in_table:
                in_table = True
                table_data = []
            table_data.append(line)
            i += 1
            continue
        elif in_table:
            # Конец таблицы
            add_table_to_doc(doc, table_data)
            table_data = []
            in_table = False
        
        # Пустые строки
        if not line.strip():
            i += 1
            continue
        
        # Заголовки
        if line.startswith('#'):
            add_heading(doc, line)
        # Жирный текст в начале строки (подзаголовки)
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip().strip('*')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
        # Обычный текст
        else:
            add_paragraph(doc, line)
        
        i += 1
    
    # Сохраняем документ
    doc.save(docx_file)
    print(f"Документ успешно сохранен: {docx_file}")

def setup_styles(doc):
    """Настраивает стили документа"""
    styles = doc.styles
    
    # Стиль для кода
    try:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_font = code_style.font
        code_font.name = 'Courier New'
        code_font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.5)
    except:
        pass  # Стиль уже существует

def add_heading(doc, line):
    """Добавляет заголовок"""
    # Подсчитываем уровень заголовка
    level = 0
    for char in line:
        if char == '#':
            level += 1
        else:
            break
    
    # Извлекаем текст заголовка
    text = line.lstrip('#').strip()
    
    # Удаляем HTML теги и ссылки
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    text = text.strip('*').strip()
    
    if text:
        heading = doc.add_heading(text, level=min(level, 3))


def add_paragraph(doc, line):
    """Добавляет параграф с форматированием"""
    # Удаляем HTML теги
    line = re.sub(r'<[^>]+>', '', line)
    
    p = doc.add_paragraph()
    
    # Разбиваем на части с форматированием
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', line)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            # Жирный текст
            run = p.add_run(part.strip('*'))
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            # Код
            run = p.add_run(part.strip('`'))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # Обычный текст
            p.add_run(part)

def add_table_to_doc(doc, table_data):
    """Добавляет таблицу в документ"""
    if len(table_data) < 2:
        return
    
    # Парсим строки таблицы
    rows = []
    for line in table_data:
        if line.strip().startswith('|'):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            # Пропускаем разделительные строки
            if cells and not all(c.replace('-', '').replace(':', '').strip() == '' for c in cells):
                rows.append(cells)
    
    if not rows:
        return
    
    # Создаем таблицу
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'
    
    # Заполняем таблицу
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            if j < len(table.rows[i].cells):
                table.rows[i].cells[j].text = cell_data

if __name__ == '__main__':
    md_file = 'docs/TZ_NowPlay_FINAL.md'
    docx_file = 'docs/TZ_NowPlay_FINAL.docx'
    
    print(f"Конвертация {md_file} в {docx_file}...")
    parse_markdown_to_docx(md_file, docx_file)
